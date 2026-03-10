from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

def create_report():
    # Load content from markdown
    markdown_path = r"c:\Users\MMohammad\Desktop\Chain-Cred-final\ChainCred_Master_Report_5000Words.md"
    with open(markdown_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Define Colors
    BNB_GOLD = RGBColor(0xF3, 0xBA, 0x2F)
    BLUE_DARK = RGBColor(0x1E, 0x3A, 0x8A)

    # 1. Title Page
    # Add Banner (if exists)
    banner_path = r"C:\Users\MMohammad\.gemini\antigravity\brain\48cc3838-dbf8-4e6c-a5ec-3a471c13edd2\chaincred_banner_1773076291480.png"
    if os.path.exists(banner_path):
        doc.add_picture(banner_path, width=Inches(6))
    
    doc.add_paragraph("\n")
    title = doc.add_heading('ChainCred: The Definitive Decentralized Credentialing Ecosystem', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph('Engineering, Strategy, and Operational Report')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = BLUE_DARK

    doc.add_page_break()

    # 2. Process the Markdown Content
    lines = content.split('\n')
    for line in lines:
        if line.startswith('###'):
            h = doc.add_heading(line.replace('###', '').strip(), 2)
            h.runs[0].font.color.rgb = BNB_GOLD
        elif line.startswith('##'):
            h = doc.add_heading(line.replace('##', '').strip(), 1)
            h.runs[0].font.color.rgb = BLUE_DARK
        elif line.startswith('#'):
            h = doc.add_heading(line.replace('#', '').strip(), 0)
            h.runs[0].font.color.rgb = BLUE_DARK
        elif line.startswith('*   '):
            p = doc.add_paragraph(line.replace('*   ', '').strip(), style='List Bullet')
        elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.'):
             # Simple list handling
             p = doc.add_paragraph(line.strip(), style='List Number')
        elif line.strip() == '---':
             doc.add_page_break()
        elif line.startswith('```'):
             # Skip code blocks for now or add them preformatted
             continue
        elif line.strip():
             # Normal paragraph
             p = doc.add_paragraph(line.strip())
             # Bold formatting
             for run in p.runs:
                 matches = re.findall(r'\*\*(.*?)\*\*', run.text)
                 for match in matches:
                     run.text = run.text.replace(f'**{match}**', match)
                     run.bold = True

    # 3. Save the report
    output_path = r"c:\Users\MMohammad\Desktop\Chain-Cred-final\ChainCred_Final_Official_Report.docx"
    doc.save(output_path)
    print(f"Report saved to: {output_path}")

if __name__ == "__main__":
    create_report()
